from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "AI_이미지분류프로젝트_최종보고서.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        set_cell_shading(header_cells[index], "E8F1F5")

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)

    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "AI 이미지 분류 프로젝트 최종 보고서"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.runs[0].font.size = Pt(9)
    paragraph.runs[0].font.color.rgb = RGBColor(90, 100, 110)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = "GitHub: whiteVR/ai-image-classification-app"
    paragraph.runs[0].font.size = Pt(9)
    paragraph.runs[0].font.color.rgb = RGBColor(90, 100, 110)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in [
        ("Title", 22),
        ("Subtitle", 12),
        ("Heading 1", 16),
        ("Heading 2", 13),
        ("Heading 3", 11),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        if "Heading" in style_name or style_name == "Title":
            style.font.bold = True
            style.font.color.rgb = RGBColor(20, 80, 100)


def build_document() -> None:
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    add_header_footer(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI 이미지 분류 프로젝트 최종 보고서")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("MNIST 숫자 분류, Cats/Dogs 분류, 웹 앱 구현, GitHub 관리, Render 및 Streamlit 배포")

    document.add_paragraph()
    add_table(
        document,
        ["항목", "내용"],
        [
            ["프로젝트 폴더", r"C:\Users\white\Documents\Codex\2026-05-09\ai-0-9-ai"],
            ["GitHub 저장소", "https://github.com/whiteVR/ai-image-classification-app"],
            ["Render 배포 URL", "https://ai-image-classification-app-iyxz.onrender.com"],
            ["주요 기술", "Python, TensorFlow/Keras, FastAPI, React TypeScript, Vite, Streamlit"],
            ["작성일", "2026-05-20"],
        ],
    )

    document.add_heading("1. 프로젝트 개요", level=1)
    document.add_paragraph(
        "이 프로젝트는 AI 기초 교육을 목적으로 이미지 분류의 전체 흐름을 단계적으로 구현한 결과물이다. "
        "처음에는 Keras의 MNIST 데이터셋을 이용해 손글씨 숫자 0부터 9까지를 CNN으로 학습하고 예측했다. "
        "이후 Kaggle Cats and Dogs 데이터셋을 이용한 CNN 분류로 확장했으며, 분류 성능 문제를 개선하기 위해 "
        "TensorFlow Datasets cats_vs_dogs와 MobileNetV2 전이학습, 실패 이미지 기반 fine-tuning까지 진행했다."
    )

    add_bullets(
        document,
        [
            "MNIST 숫자 분류와 직접 작성한 숫자 이미지 예측 구현",
            "Kaggle Cats and Dogs 데이터셋 기반 CNN 학습과 예측 구현",
            "TensorFlow Datasets cats_vs_dogs와 MobileNetV2 전이학습 적용",
            "실패 이미지 하드케이스를 추가하여 MobileNetV2 상위 레이어 fine-tuning 진행",
            "FastAPI + React TypeScript 앱 구현 및 Render 배포",
            "Streamlit 버전 앱 추가 및 무료 배포 대안 문서화",
            "GitHub 저장소 기반 버전 관리와 배포 자동화 흐름 구성",
        ],
    )

    document.add_heading("2. 교육 과정 문서화", level=1)
    add_table(
        document,
        ["차시", "주제", "산출 문서"],
        [
            ["1차시", "AI, 머신러닝, 딥러닝, 이미지 픽셀, 학습/테스트 데이터", "docs/01_ai_image_classification_concepts.md"],
            ["2차시", "MNIST 데이터셋, CNN 학습, 정확도/손실 그래프, 틀린 예측 분석", "docs/02_mnist_digit_classification.md"],
            ["3차시", "직접 쓴 숫자 이미지 예측, 전처리, 확률 해석", "docs/03_predict_custom_digit_image.md"],
            ["4차시", "Cats/Dogs 분류, 실제 사진 데이터, 과적합, 데이터 증강", "docs/04_cats_vs_dogs_classification.md"],
            ["5단계", "FastAPI + React 앱 구성, Git 관리, 배포 절차", "docs/05_app_solution_and_deployment.md"],
            ["6단계", "Render 배포 가이드와 장애 해결", "docs/06_render_deployment_guide.md"],
            ["7단계", "Streamlit 앱과 무료 배포 대안", "docs/07_streamlit_deployment_guide.md"],
        ],
    )

    document.add_heading("3. 모델 학습과 예측 산출물", level=1)
    add_table(
        document,
        ["구분", "학습 방식", "모델/결과물 위치"],
        [
            ["MNIST", "Keras MNIST + CNN", "models01/mnist_digit_cnn.keras, output01, predict_outputs01"],
            ["Cats/Dogs 1차", "Kaggle Cats and Dogs + CNN", "models_CD/cats_vs_dogs_cnn.keras, output_CD"],
            ["Cats/Dogs 2차", "TFDS cats_vs_dogs + MobileNetV2 전이학습", "models_TFDS_CD, output_TFDS_CD"],
            ["Cats/Dogs 3차", "실패 이미지 추가 + MobileNetV2 fine-tuning", "models_TFDS_CD_FT, output_TFDS_CD_FT"],
        ],
    )

    document.add_heading("4. 앱 아키텍처", level=1)
    document.add_paragraph(
        "최종 앱은 Keras 모델 추론을 안정적으로 수행하기 위해 Python 백엔드를 유지하고, 사용자 화면은 React TypeScript로 구성했다. "
        "React만으로 브라우저에서 Keras 모델을 직접 실행하려면 TensorFlow.js 변환이 필요하므로, 현재 산출물에는 FastAPI 백엔드와 "
        "React 프론트엔드를 분리한 구조가 가장 적합하다."
    )
    add_table(
        document,
        ["영역", "기술", "역할"],
        [
            ["백엔드", "FastAPI", "이미지 업로드 API, Keras 모델 로딩, MNIST 및 Cats/Dogs 예측"],
            ["프론트엔드", "React TypeScript + Vite", "이미지 업로드, 모델 선택, 확률 시각화"],
            ["모델", "TensorFlow/Keras", "CNN, MobileNetV2, fine-tuned MobileNetV2 추론"],
            ["배포", "Render", "FastAPI가 React dist 정적 파일과 API를 함께 제공"],
            ["대안 앱", "Streamlit", "Python 단일 파일 기반 AI 데모 앱"],
        ],
    )

    document.add_heading("5. 주요 앱 파일", level=1)
    add_table(
        document,
        ["파일", "설명"],
        [
            ["app/backend/inference.py", "모델 로딩, 이미지 전처리, MNIST 및 Cats/Dogs 예측 함수"],
            ["app/backend/main.py", "FastAPI 엔드포인트와 React 정적 파일 서빙"],
            ["app/frontend/src/App.tsx", "React TypeScript 기반 웹 UI"],
            ["app/frontend/src/styles.css", "앱 화면 스타일"],
            ["streamlit_app.py", "Streamlit 버전 이미지 분류 앱"],
            ["render.yaml", "Render 배포 설정"],
            [".python-version", "Render Python 3.12.7 고정"],
        ],
    )

    document.add_heading("6. GitHub 버전 관리", level=1)
    add_table(
        document,
        ["커밋", "내용"],
        [
            ["e7d8c13", "FastAPI + React 앱, 모델/결과물/문서 최초 커밋"],
            ["a0fc1ce", "MNIST 문서 수정과 숫자 샘플 추가"],
            ["efda43b", "Render 배포 설정 추가"],
            ["ddbe25c", "Render 무료 플랜 설정"],
            ["8270447", "Render 빌드 설정 수정"],
            ["7a423fb", "Python 3.12.7 버전 고정"],
            ["ff6b220", "Streamlit 이미지 분류 앱 추가"],
        ],
    )

    document.add_heading("7. Render 배포와 문제 해결", level=1)
    document.add_paragraph(
        "Render 배포 과정에서는 포트 점유, 결제 정보 요구, uvicorn 미설치, Python 3.14와 TensorFlow 호환성 문제가 순서대로 발생했다. "
        "각 문제는 로그를 기준으로 원인을 분리하고 설정을 조정해 해결했다."
    )
    add_table(
        document,
        ["문제", "원인", "조치"],
        [
            ["WinError 10013", "8000번 포트가 이미 로컬 Python 프로세스에서 사용 중", "기존 백엔드 동작 확인 후 프론트엔드만 실행"],
            ["Payment Information Required", "render.yaml의 plan이 starter로 설정됨", "plan: free로 변경"],
            ["No module named uvicorn", "Render 빌드 단계에서 requirements 설치가 반영되지 않음", "buildCommand 단순화 및 dist 포함"],
            ["No matching distribution for tensorflow", "Render가 Python 3.14로 빌드", ".python-version에 3.12.7 지정"],
            ["CUDA driver warning", "Render 무료 서버에 GPU 없음", "CPU 실행 경고로 판단, 조치 불필요"],
        ],
    )

    document.add_heading("8. 실행 방법", level=1)
    document.add_heading("8.1 FastAPI + React 로컬 실행", level=2)
    add_code_block(
        document,
        [
            r"cd C:\Users\white\Documents\Codex\2026-05-09\ai-0-9-ai",
            r".\aivenv\Scripts\Activate.ps1",
            r"python -m uvicorn app.backend.main:app --reload --host 127.0.0.1 --port 8000",
            r"cd app\frontend",
            r"npm run dev",
        ],
    )
    document.add_paragraph("접속 주소: http://localhost:5173")

    document.add_heading("8.2 Render 배포 앱", level=2)
    add_code_block(
        document,
        [
            "https://ai-image-classification-app-iyxz.onrender.com",
            "https://ai-image-classification-app-iyxz.onrender.com/api/health",
        ],
    )

    document.add_heading("8.3 Streamlit 로컬 실행", level=2)
    add_code_block(
        document,
        [
            r"cd C:\Users\white\Documents\Codex\2026-05-09\ai-0-9-ai",
            r".\aivenv\Scripts\Activate.ps1",
            r"pip install -r requirements-streamlit.txt",
            r"streamlit run streamlit_app.py",
        ],
    )
    document.add_paragraph("접속 주소: http://localhost:8501")

    document.add_heading("9. 배포 전략 비교", level=1)
    add_table(
        document,
        ["플랫폼", "적합도", "비고"],
        [
            ["Render", "현재 배포 완료", "FastAPI + React + Keras 모델을 단일 서비스로 운영"],
            ["Streamlit Community Cloud", "대안으로 적합", "Python 단일 앱 구조, 교육용 데모에 적합"],
            ["Hugging Face Spaces", "AI 데모에 매우 적합", "Streamlit 또는 Docker Space로 배포 가능"],
            ["Vercel/Netlify", "프론트엔드만 적합", "Keras Python 추론 백엔드가 별도로 필요"],
        ],
    )

    document.add_heading("10. 최종 상태", level=1)
    add_bullets(
        document,
        [
            "MNIST 숫자 분류와 예측이 정상 동작한다.",
            "Cats/Dogs fine-tuned MobileNetV2 모델 예측이 정상 동작한다.",
            "FastAPI + React 앱이 Render에서 정상 배포되어 /api/health 응답을 반환한다.",
            "React 앱 화면에서 MNIST 이미지 업로드와 예측 결과 표시가 확인되었다.",
            "Streamlit 버전 앱이 로컬에서 정상 응답한다.",
            "GitHub 저장소에 코드, 문서, 모델, 결과물이 커밋되어 있다.",
        ],
    )

    document.add_paragraph()
    document.add_paragraph(
        "결론적으로 이 프로젝트는 AI 이미지 분류 교육 콘텐츠, 모델 학습 코드, 예측 코드, 웹 앱, 대체 Streamlit 앱, "
        "GitHub 버전 관리, Render 배포까지 포함한 end-to-end 실습 프로젝트로 정리되었다."
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
